"""Self Learning Material (SLM) as editable DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from .models import Subject, UnitContentPackage


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def generate_slm_docx(
    output_path: Path,
    *,
    program_name: str,
    subject: Subject,
    unit_label: str,
    pack: UnitContentPackage,
) -> Path:
    doc = Document()
    _heading(doc, program_name, 0)
    _heading(doc, subject.name, 1)
    _heading(doc, f"Unit {unit_label}: {pack.unit_title}", 2)

    meta = doc.add_paragraph()
    meta.add_run(f"Credits: {subject.credits} | Target SLM words: ~{pack.words_target_slm:,} | Source: {pack.generated_by}")

    _heading(doc, "1. Introduction", 2)
    _para(doc, pack.introduction)

    _heading(doc, "2. Learning Objectives (Bloom's Taxonomy)", 2)
    _bullets(doc, pack.learning_objectives)

    _heading(doc, "3. Unit Content — Topics", 2)
    for i, slide in enumerate(pack.topic_slides, 1):
        _heading(doc, f"3.{i} {slide.title}", 3)
        _bullets(doc, slide.bullets)
        if slide.speaker_notes:
            _para(doc, slide.speaker_notes)

    _heading(doc, "4. Case Study", 2)
    _para(doc, pack.case_study)

    _heading(doc, "5. Summary", 2)
    _para(doc, pack.summary)

    _heading(doc, "6. Keywords / Glossary", 2)
    _bullets(doc, pack.keywords)

    _heading(doc, "7. Check Your Progress", 2)
    _bullets(doc, pack.check_your_progress)

    _heading(doc, "8. Self-Assessment Questions (SAQs)", 2)
    for i, q in enumerate(pack.saqs, 1):
        doc.add_paragraph(f"Q{i}. {q}", style="List Number")

    _heading(doc, "9. Answers to SAQs", 2)
    for i, a in enumerate(pack.saq_answers, 1):
        _para(doc, f"A{i}. {a}")

    if subject.course_outcomes:
        _heading(doc, "10. Course Outcomes (Reference)", 2)
        _bullets(doc, subject.course_outcomes[:6])

    _heading(doc, "11. References", 2)
    _bullets(doc, pack.references)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
