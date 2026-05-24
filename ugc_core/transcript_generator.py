"""Faculty delivery transcript (video-ready script)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .models import Subject, UnitContentPackage


def generate_transcript_docx(
    output_path: Path,
    *,
    program_name: str,
    subject: Subject,
    unit_label: str,
    pack: UnitContentPackage,
    delivered_by: str = "",
) -> Path:
    doc = Document()
    doc.add_heading(f"{subject.name} — Unit {unit_label}", 0)
    doc.add_heading("e-Tutorial Transcript (Faculty Delivery)", level=1)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Program: {program_name} | Duration: ~{pack.minutes_per_unit} minutes | "
        f"Slides: ~{pack.target_slide_count} | Faculty: {delivered_by or 'TBD'}"
    )

    doc.add_heading("Session overview", level=2)
    doc.add_paragraph(pack.introduction)

    doc.add_heading("Full narration script", level=2)
    body = doc.add_paragraph(pack.transcript)
    body.style.font.size = Pt(11)

    doc.add_heading("Slide-wise cues", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Slide"
    hdr[1].text = "Title"
    hdr[2].text = "Teaching note"

    for i, slide in enumerate(pack.topic_slides, 1):
        row = table.add_row().cells
        row[0].text = str(i + 4)
        row[1].text = slide.title
        row[2].text = slide.speaker_notes or slide.visual_suggestion

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
