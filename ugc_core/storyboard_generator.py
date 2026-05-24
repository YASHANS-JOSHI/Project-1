"""Slide storyboard for instructional design / video production."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .models import UnitContentPackage


def generate_storyboard_docx(
    output_path: Path,
    *,
    pack: UnitContentPackage,
) -> Path:
    doc = Document()
    doc.add_heading(f"Storyboard — {pack.unit_title}", 0)

    rows = [
        ("1", "Title", pack.unit_title, "Show course branding", "Open with greeting"),
        ("2", "Syllabus", "Unit overview", "Text + icons", pack.introduction[:200]),
        ("3", "Topics", "Topic map", "Bullet list layout", "Preview session flow"),
        ("4", "Objectives", "Learning Objectives", "Bloom verbs highlighted", "State measurable outcomes"),
    ]
    for i, slide in enumerate(pack.topic_slides, 1):
        rows.append(
            (
                str(4 + i),
                "Content",
                slide.title,
                slide.visual_suggestion or "Relevant diagram",
                slide.speaker_notes[:300] if slide.speaker_notes else "Explain bullets; pause for questions",
            )
        )
    n = 5 + len(pack.topic_slides)
    rows.append((str(n), "Summary", "Key takeaways", "Recap graphic", pack.summary[:200]))
    rows.append((str(n + 1), "Closing", "Happy Learning", "Brand slide", "Sign off"))

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    h = table.rows[0].cells
    for i, label in enumerate(["#", "Type", "On-screen text", "Visual", "Faculty instruction"]):
        h[i].text = label

    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
